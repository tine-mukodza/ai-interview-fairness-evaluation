
import streamlit as st
import pandas as pd
from datetime import datetime
import os
import re
import random
import io

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# ---------------------------------------------------------
# PAGE CONFIG + GLOBAL CSS
# ---------------------------------------------------------
st.set_page_config(
    page_title="AI Interview Transparency (HCI Case Study)",
    page_icon="🧠",
    layout="wide"
)

st.markdown(
    """
    <style>
    div.block-container {
        max-width: 1100px;
        padding-top: 1.2rem;
    }
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}

    /* General typography */
    h1, h2, h3 {
        font-family: system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
    }
    .stMarkdown h3 {
        font-size: 1.05rem !important;
        margin-bottom: 0.35rem;
    }

    /* Make expander headers a bit stronger */
    .streamlit-expanderHeader {
        font-weight: 650 !important;
    }

    /* Step dots (non-clickable, no "button" affordance) */
    .step-wrap {
        display:flex;
        align-items:center;
        gap:10px;
        margin: 0.35rem 0 0.25rem 0;
        user-select:none;
    }
    .step-dot {
        width: 10px;
        height: 10px;
        border-radius: 999px;
        border: 2px solid #64748b;
        background: transparent;
        opacity: 0.75;
    }
    .step-dot-done {
        border-color: #22c55e;
        background: #22c55e;
        opacity: 1;
    }
    .step-dot-active {
        border-color: #3b82f6;
        background: #3b82f6;
        opacity: 1;
        box-shadow: 0 0 0 4px rgba(59,130,246,0.15);
    }
    .step-line {
        flex: 1;
        height: 2px;
        background: #334155;
        opacity: 0.6;
        border-radius: 999px;
    }
    .step-labels {
        display:flex;
        justify-content:space-between;
        font-size: 0.78rem;
        color: #94a3b8;
        margin-bottom: 0.6rem;
        user-select:none;
    }

    /* Prevent "hand" cursor on custom elements */
    .step-wrap, .step-labels { cursor: default !important; }
    
    /* Select-like display (non-editable, looks like dropdown) */
    .select-like {
        padding: 0.55rem 0.75rem;
        border-radius: 0.5rem;
        border: 1px solid rgba(148,163,184,0.35);
        background: rgba(2,6,23,0.85);
        color: #e5e7eb;
        font-size: 0.95rem;
        line-height: 1.35;
        user-select: text;
    }

    </style>
    """,
    unsafe_allow_html=True,
)

# ---------------------------------------------------------
# HELPERS & CORE LOGIC
# ---------------------------------------------------------
def card(title, body="", icon=""):
    """Simple bordered card block with optional title."""
    with st.container(border=True):
        if title:
            st.markdown(f"### {icon} {title}")
        if body:
            st.write(body)

def extract_keywords(text: str):
    """Very lightweight keyword extractor (no ML)."""
    words = re.findall(r"[A-Za-z']+", text.lower())
    stop = set(
        [
            "the","a","an","and","or","to","of","in","on","for","with",
            "my","your","our","their","you","i","we","was","were","is",
            "are","that","this","from","have","has","had","been","at",
            "as","by","it","itself",
        ]
    )
    keywords = [w for w in words if w not in stop and len(w) > 3]
    return list(dict.fromkeys(keywords))[:8]

VALUES = {
    "Collaboration": ["team", "together", "support", "conflict", "help"],
    "Integrity": ["ethical", "honest", "truth", "responsible", "fair"],
    "Ownership": ["initiative", "led", "managed", "owned", "accountable"],
    "Customer Focus": ["customer", "client", "user", "service", "needs"],
    "Data Responsibility": ["data", "privacy", "security", "accuracy", "bias"],
}

SCENARIOS = [
    {
        "name": "Scenario 1 – Collaboration (Team Conflict)",
        "prompt": "Tell me about a time you worked with a team to solve a difficult problem?",
        "value": "Collaboration",
    },
    {
        "name": "Scenario 2 – Integrity (Ethical Dilemma)",
        "prompt": "Describe a situation where you had to choose the ethical option under pressure?",
        "value": "Integrity",
    },
    {
        "name": "Scenario 3 – Ownership (Taking Initiative)",
        "prompt": "Tell me about a time you took initiative without being asked?",
        "value": "Ownership",
    },
    {
        "name": "Scenario 4 – Data Responsibility (Handling Sensitive Info)",
        "prompt": "Describe a moment when you handled sensitive data or ensured data accuracy?",
        "value": "Data Responsibility",
    },
    {
        "name": "Scenario 5 – Customer Focus (User Impact)",
        "prompt": "Tell me about a time you improved a customer or user experience?",
        "value": "Customer Focus",
    },
]

FOLLOWUP_BANK = {
    "Collaboration": [
        "What role did you personally play in helping the team succeed?",
        "How did you handle disagreement or tension in the group?",
        "What did you learn about teamwork from that experience?",
    ],
    "Integrity": [
        "What made that decision ethically difficult?",
        "How did you communicate your choice to others?",
        "Looking back, would you do anything differently?",
    ],
    "Ownership": [
        "What motivated you to take initiative in that situation?",
        "How did you measure success for that project?",
        "What obstacles did you face and how did you handle them?",
    ],
    "Customer Focus": [
        "How did you identify what the customer or user actually needed?",
        "What change did you make and what was its impact?",
        "How did you gather feedback after your solution?",
    ],
    "Data Responsibility": [
        "How did you make sure the data was accurate or handled safely?",
        "What risks did you consider when working with that data?",
        "How did your actions protect stakeholders or users?",
    ],
}

def detect_value_tag(answer_text: str):
    """Heuristic value detector based on word matches."""
    answer_text = answer_text.lower()
    scores = {v: 0 for v in VALUES.keys()}
    for v, kws in VALUES.items():
        for k in kws:
            if k in answer_text:
                scores[v] += 1

    best_value = max(scores, key=scores.get)
    if scores[best_value] == 0:
        return "General Professionalism", "Low"
    elif scores[best_value] <= 2:
        return best_value, "Medium"
    else:
        return best_value, "High"

def generate_followup(resume_text: str, answer_text: str, chosen_value: str):
    """Generates a follow-up question + explanation using earlier logic."""
    detected_value, confidence = detect_value_tag(answer_text)

    # Respect the scenario target value; report detected as internal guess only
    value_tag = chosen_value
    followup = random.choice(FOLLOWUP_BANK[value_tag])

    resume_kws = extract_keywords(resume_text)
    answer_kws = extract_keywords(answer_text)

    reasoning = (
        f"The follow-up targets **{value_tag}** based on the scenario you selected. "
        f"In your resume, I noticed: {', '.join(resume_kws) or 'no clear keywords'}. "
        f"In your answer, I noticed: {', '.join(answer_kws) or 'no clear keywords'}. "
        f"The system's internal guess from your answer alone was **{detected_value}** "
        f"with **{confidence}** confidence."
    )

    return followup, reasoning, value_tag, confidence, resume_kws, answer_kws

def neutralize_question(q: str) -> str:
    """Softens a question if the participant flags it as unfair."""
    if not q:
        return ""
    return "In any context you’re comfortable sharing, " + q[0].lower() + q[1:]


def generate_alternative_followup(current_followup: str, value_tag: str) -> str:
    """Return a different follow-up from the same value bank (simple alternative)."""
    bank = FOLLOWUP_BANK.get(value_tag, [])
    if not bank:
        return ""
    # Prefer an option different from the current follow-up
    options = [q for q in bank if q != current_followup] or bank
    return random.choice(options)


# ---------------------------------------------------------
# LOGGING
# ---------------------------------------------------------
LOG_FILE = "interview_logs.csv"

def log_row(row: dict):
    df_row = pd.DataFrame([row])
    if os.path.exists(LOG_FILE):
        df_existing = pd.read_csv(LOG_FILE)
        df_all = pd.concat([df_existing, df_row], ignore_index=True)
    else:
        df_all = df_row
    df_all.to_csv(LOG_FILE, index=False)

# ---------------------------------------------------------
# EXPORT HELPERS (CSV / Excel / Word / PDF)
# ---------------------------------------------------------
def df_to_excel_bytes(df: pd.DataFrame) -> bytes:
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="logs")
    return bio.getvalue()


def row_to_word_bytes(row: dict) -> bytes:
    """
    Create a readable, single-session Word summary for participants / research appendix.

    Formatting goal: clear key-value lines with bold labels and numbered responses where applicable.
    """
    doc = Document()

    def add_kv(label: str, value: str):
        p = doc.add_paragraph()
        run = p.add_run(f"{label}")
        run.bold = True
        p.add_run(f" {value if value is not None else ''}")

    def add_numbered_response(text_value: str):
        lines = [ln.strip() for ln in (text_value or "").splitlines() if ln.strip()]
        if not lines:
            doc.add_paragraph("(no response provided)")
            return
        if len(lines) == 1:
            doc.add_paragraph(lines[0])
            return
        for ln in lines:
            doc.add_paragraph(ln, style="List Number")

    # Title
    title = doc.add_paragraph()
    r = title.add_run("AI Interview Prototype – Session Summary")
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph("")

    # Metadata
    add_kv("Timestamp:", row.get("timestamp", ""))
    add_kv("Participant ID:", row.get("participant_id", ""))
    add_kv("Scenario:", row.get("scenario", ""))
    add_kv("Target value:", row.get("target_value", ""))

    doc.add_paragraph("")
    doc.add_paragraph("Inputs").runs[0].bold = True

    add_kv("Resume summary / text:", "")
    resume_txt = (row.get("resume_text", "") or "").strip()
    doc.add_paragraph(resume_txt if resume_txt else "(not provided)")

    doc.add_paragraph("")
    add_kv("Scenario Question:", "")
    doc.add_paragraph((row.get("scenario_prompt_used", "") or "").strip() or "(not provided)")

    doc.add_paragraph("")
    add_kv("Scenario answer:", "")
    add_numbered_response(row.get("answer_text", ""))

    doc.add_paragraph("")
    doc.add_paragraph("AI Follow-Up").runs[0].bold = True

    add_kv("Follow-up question:", (row.get("followup_question", "") or "").strip())
    doc.add_paragraph("")
    add_kv("Follow-up answer:", "")
    add_numbered_response(row.get("followup_answer_text", ""))

    doc.add_paragraph("")
    add_kv("Value tag:", row.get("value_tag", ""))
    add_kv("Confidence:", row.get("confidence", ""))

    doc.add_paragraph("")
    add_kv("Reasoning summary:", "")
    doc.add_paragraph((row.get("reasoning_summary", "") or "").strip() or "(not provided)")

    doc.add_paragraph("")
    doc.add_paragraph("Fairness / Contestability").runs[0].bold = True
    add_kv("Flagged as unfair / uncomfortable:", str(row.get("flag_unfair", "")))

    neutral = (row.get("neutralized_question", "") or "").strip()
    if neutral:
        add_kv("In any context you’re comfortable sharing:", neutral)

    doc.add_paragraph("")
    add_kv("Optional: What felt unfair or uncomfortable ?", "")
    unfair = (row.get("unfair_comment", "") or "").strip()
    doc.add_paragraph(unfair if unfair else "(not provided)")

    alt_q = (row.get("alternative_question", "") or "").strip()
    if alt_q:
        doc.add_paragraph("")
        add_kv("Optional: Alternative question:", alt_q)

        doc.add_paragraph("")
        add_kv("Alternative answer:", "")
        add_numbered_response(row.get("alternative_answer_text", ""))

    doc.add_paragraph("")
    doc.add_paragraph("Ratings and Feedback").runs[0].bold = True
    add_kv("Fairness score:", str(row.get("fairness_score", "")))
    add_kv("Relevance score:", str(row.get("relevance_score", "")))
    add_kv("Comfort score:", str(row.get("comfort_score", "")))
    add_kv("Trust score:", str(row.get("trust_score", "")))
    add_kv("Accept AI:", str(row.get("accept_ai", "")))

    feedback = (row.get("open_feedback", "") or "").strip()
    if feedback:
        doc.add_paragraph("")
        add_kv("Open feedback:", "")
        doc.add_paragraph(feedback)

    doc.add_paragraph("")
    note = doc.add_paragraph()
    note_run = note.add_run("Notes: ")
    note_run.bold = True
    note.add_run("Open feedback and full texts (resume/answer) are saved in the CSV/Excel downloads.")

    # Styling
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.size is None:
                run.font.size = Pt(11)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def row_to_pdf_bytes(row: dict) -> bytes:
    """
    Create a readable single-session PDF summary with clear labels.
    """
    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=letter)
    width, height = letter

    x = 50
    y = height - 55
    line_h = 14

    def ensure_space(lines_needed=1):
        nonlocal y
        if y < 70 + (lines_needed * line_h):
            c.showPage()
            y = height - 55

    def draw_label_value(label: str, value: str, bold_label=True):
        nonlocal y
        ensure_space(1)
        c.setFont("Helvetica-Bold" if bold_label else "Helvetica", 10)
        c.drawString(x, y, label)
        c.setFont("Helvetica", 10)
        # wrap value
        value = value if value is not None else ""
        max_chars = 95
        parts = [value[i:i+max_chars] for i in range(0, len(value), max_chars)] or [""]
        if parts:
            c.drawString(x + 140, y, parts[0])
        y -= line_h
        for part in parts[1:]:
            ensure_space(1)
            c.drawString(x + 140, y, part)
            y -= line_h

    # Title
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x, y, "AI Interview Prototype – Session Summary")
    y -= 24

    # Metadata
    draw_label_value("Timestamp:", str(row.get("timestamp", "")))
    draw_label_value("Participant ID:", str(row.get("participant_id", "")))
    draw_label_value("Scenario:", str(row.get("scenario", "")))
    draw_label_value("Target value:", str(row.get("target_value", "")))
    y -= 8

    # Follow-up
    c.setFont("Helvetica-Bold", 12)
    c.drawString(x, y, "AI Follow-Up")
    y -= 18

    draw_label_value("Follow-up question:", str(row.get("followup_question", "")))
    draw_label_value("Follow-up answer:", str(row.get("followup_answer_text", "")))
    draw_label_value("Value tag:", str(row.get("value_tag", "")))
    draw_label_value("Confidence:", str(row.get("confidence", "")))
    y -= 8

    # Ratings (participant feedback)
    c.setFont("Helvetica-Bold", 12)
    c.drawString(x, y, "Ratings")
    y -= 18
    draw_label_value("Fairness score:", str(row.get("fairness_score", "")))
    draw_label_value("Relevance score:", str(row.get("relevance_score", "")))
    draw_label_value("Comfort score:", str(row.get("comfort_score", "")))
    draw_label_value("Trust score:", str(row.get("trust_score", "")))
    draw_label_value("Accept AI:", str(row.get("accept_ai", "")))
    draw_label_value("Flag unfair:", str(row.get("flag_unfair", "")))
    y -= 8

    # Fairness / alternative
    c.setFont("Helvetica-Bold", 12)
    c.drawString(x, y, "Fairness / Contestability")
    y -= 18

    draw_label_value("Flagged:", str(row.get("flag_unfair", "")))
    neutral = str(row.get("neutralized_question", "") or "")
    if neutral.strip():
        draw_label_value("In any context you’re comfortable sharing:", neutral)

    unfair = str(row.get("unfair_comment", "") or "")
    if unfair.strip():
        draw_label_value("What felt unfair or uncomfortable?:", unfair)

    alt_q = str(row.get("alternative_question", "") or "")
    if alt_q.strip():
        draw_label_value("Alternative question:", alt_q)
        draw_label_value("Alternative answer:", str(row.get("alternative_answer_text", "")))

    y -= 8
    c.setFont("Helvetica-Bold", 10)
    ensure_space(2)
    c.drawString(x, y, "Notes:")
    c.setFont("Helvetica", 10)
    c.drawString(x + 45, y, "Open feedback and full texts (resume/answer) are saved in the CSV/Excel downloads.")
    y -= line_h

    c.save()
    return bio.getvalue()

# ---------------------------------------------------------
# SESSION STATE (4 steps after consent)
# ---------------------------------------------------------
defaults = {
    "consent": False,

    # progress flags
    "resume_done": False,
    "scenario_answer_done": False,
    "followup_answer_text": "",
    "unfair_details": "",
    "alternative_question": "",
    "alternative_answer_text": "",
    "followup_generated": False,
    "followup_response_done": False,
    "followup_done": False,

    # step control
    "active_step": 1,  # 1–4 after consent
    "exp1_open": True,
    "exp2_open": False,
    "exp3_open": False,
    "exp4_open": False,

    # scenario state
    "scenario_prompt": "",
    "scenario_name": SCENARIOS[0]["name"],
    "prev_scenario_name": None,
}

for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ---------------------------------------------------------
# UI HELPERS – NON-CLICKY STEP INDICATOR
# ---------------------------------------------------------
def render_step_dots():
    labels = ["Resume", "Pick a scenario", "AI follow‑up", "Feedback"]
    active = st.session_state.active_step

    # done logic
    done = {
        1: st.session_state.resume_done,
        2: st.session_state.scenario_answer_done,
        3: st.session_state.followup_response_done,
        4: False,  # feedback is "done" only after save; not tracked separately
    }

    dots_html = ["<div class='step-wrap'>"]
    for i in range(1, 5):
        cls = "step-dot"
        if i == active:
            cls += " step-dot-active"
        elif done.get(i, False):
            cls += " step-dot-done"
        dots_html.append(f"<div class='{cls}'></div>")
        if i != 4:
            dots_html.append("<div class='step-line'></div>")
    dots_html.append("</div>")

    st.markdown("".join(dots_html), unsafe_allow_html=True)

    # labels line
    st.markdown(
        f"""
        <div class='step-labels'>
            <div>1. {labels[0]}</div>
            <div>2. {labels[1]}</div>
            <div>3. {labels[2]}</div>
            <div>4. {labels[3]}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

# ---------------------------------------------------------
# SIDEBAR
# ---------------------------------------------------------
with st.sidebar:
    st.header("Study Settings")
    st.text_input("Participant ID (optional)", key="participant_id", placeholder="P01, P02 …")
    st.markdown("---")
    st.caption("Target values in this study:")
    st.write(", ".join(VALUES.keys()))
    st.markdown("---")
    if st.button("Reset session"):
        st.session_state.clear()
        st.rerun()

    st.markdown("---")
    st.subheader("Researcher view (optional)")
    admin_password = os.environ.get("ADMIN_PASSWORD", "")
    if admin_password:
        entered = st.text_input("Admin password", type="password", key="admin_pw")
        if entered == admin_password:
            if os.path.exists(LOG_FILE):
                df_logs = pd.read_csv(LOG_FILE)
                st.metric("Total submissions", len(df_logs))
                if "participant_id" in df_logs.columns:
                    non_empty = df_logs["participant_id"].fillna("").astype(str).str.strip()
                    st.metric("Unique participant IDs", non_empty[non_empty != ""].nunique())
                if "scenario" in df_logs.columns:
                    st.caption("Scenario counts")
                    st.dataframe(df_logs["scenario"].value_counts().rename_axis("scenario").reset_index(name="count"), use_container_width=True)

                with open(LOG_FILE, "rb") as f:
                    st.download_button("Download research CSV (all sessions)", data=f, file_name="interview_logs.csv", mime="text/csv")

                st.download_button(
                    "Download research Excel (all sessions)",
                    data=df_to_excel_bytes(df_logs),
                    file_name="interview_logs.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            else:
                st.info("No submissions yet (log file not found).")
        elif entered:
            st.error("Incorrect password.")
    else:
        st.caption("To enable researcher summaries, set an environment variable ADMIN_PASSWORD on the server.")

# ---------------------------------------------------------
# HEADER + STUDY ABOUT (moved up) + PROGRESS (redesigned)
# ---------------------------------------------------------
st.title("AI Interview Transparency")
st.caption("A Human-Centered HCI Case Study")

with st.expander("What is this study about?", expanded=False):
    st.write(
        "- Test an AI interviewer that generates value-aligned follow-up questions.\n"
        "- Demonstrate transparently how the AI uses your resume text and your answer to create the follow-up questions.\n"
        "- Ask you to rate how fair, appropriate, and clear the follow-up question feels.\n"
        "- No real name is required."
    )

st.markdown(f"**Progress: Step {st.session_state.active_step} of 4**")
st.progress((st.session_state.active_step - 1) / 4)
render_step_dots()
st.caption("Only the current step is expanded. You can re-open earlier steps any time to review or edit.")

# ---------------------------------------------------------
# CONSENT GATE
# ---------------------------------------------------------
card(
    "Consent to Participate",
    icon="✅",
    body=(
        "- Participation is voluntary.\n"
        "- You may paste either a short summary or a resume.\n"
        "- We only use responses for a classroom research project on fairness and usability.\n"
        "- You can stop at any time by closing the page or using the reset button."
    ),
)

consent = st.checkbox(
    "I have read this summary and I agree to take part in this prototype study.",
    value=st.session_state.consent,
)
st.session_state.consent = consent

if not st.session_state.consent:
    st.info("Please give your consent above to start the interview steps.")
    st.stop()

# ---------------------------------------------------------
# STEP 1 – RESUME / EXPERIENCE
# ---------------------------------------------------------
step1_label = "Step 1 – Resume / Experience Context"
if st.session_state.resume_done and st.session_state.active_step != 1:
    rt_len = len(st.session_state.get("resume_text", "").strip())
    step1_label += f"  ✅ (about {rt_len} characters)"

with st.expander(step1_label, expanded=st.session_state.exp1_open):
    card(
        "",
        icon="📄",
        body=(
            "Paste a short summary or your full resume text. Personal details are optional. "
            "We only use this text so the AI can personalize its follow-up question."
        ),
    )

    st.text_area(
        "Resume or summary",
        key="resume_text",
        height=220,
        placeholder=(
            "Example:\n"
            "I have experience in SQL, Power BI, and team-based analytics projects in finance "
            "and operations. Recently I..."
        ),
    )

    if st.session_state.get("resume_text"):
        st.caption(f"Characters provided: {len(st.session_state['resume_text'])}")

    if st.button("Save and continue", key="btn_step1"):
        if not st.session_state["resume_text"].strip():
            st.error("Please add at least a short summary or resume text before continuing.")
        else:
            st.session_state.resume_done = True
            st.session_state.active_step = 2
            st.session_state.exp1_open = False
            st.session_state.exp2_open = True
            st.success("Resume saved. Moving to Step 2.")
            st.rerun()

# ---------------------------------------------------------
# STEP 2 – PICK A SCENARIO + ANSWER (save happens AFTER answer)
# ---------------------------------------------------------
if st.session_state.resume_done:
    step2_label = "Step 2 – Pick a scenario"
    if st.session_state.scenario_answer_done and st.session_state.active_step != 2:
        step2_label += "  ✅ (scenario + answer saved)"

    with st.expander(step2_label, expanded=st.session_state.exp2_open):

        scenario_names = [s["name"] for s in SCENARIOS]

        st.selectbox(
            "Pick a scenario",
            options=scenario_names,
            key="scenario_name",
        )

        # Reset prompt when scenario changes
        if st.session_state.prev_scenario_name != st.session_state.scenario_name:
            base_scenario = next(s for s in SCENARIOS if s["name"] == st.session_state.scenario_name)
            st.session_state.scenario_prompt = base_scenario["prompt"]
            st.session_state.prev_scenario_name = st.session_state.scenario_name

        # Scenario question (read-only; styled like a dropdown)
        st.markdown("**Scenario Question**")
        st.markdown(
            f"<div class='select-like'>{st.session_state.get('scenario_prompt', '')}</div>",
            unsafe_allow_html=True,
        )

        # Participant answer is part of Step 2 now
        st.markdown("**Your response**")
        st.text_area(
            "Type your response to the scenario prompt",
            key="answer_text",
            height=160,
            placeholder="Type your answer here…",
        )

        if st.button("Save scenario + response and continue", key="btn_step2"):
            if not st.session_state.get("scenario_prompt", "").strip():
                st.error("Scenario prompt is empty. Please select a scenario or enter a prompt.")
            elif not st.session_state.get("answer_text", "").strip():
                st.error("Please write your response before continuing.")
            else:
                st.session_state.scenario_answer_done = True
                st.session_state.active_step = 3
                st.session_state.exp2_open = False
                st.session_state.exp3_open = True
                st.success("Saved. Moving to Step 3.")
                st.rerun()
else:
    st.info("Complete Step 1 first. Step 2 will appear after you save your resume/experience.")

# ---------------------------------------------------------
# STEP 3 – AI FOLLOW-UP + EXPLAINABILITY (was Step 4)
# ---------------------------------------------------------
if st.session_state.scenario_answer_done:
    step3_label = "Step 3 – AI Follow-Up Question & Explanation"
    if st.session_state.followup_done and st.session_state.active_step != 3:
        step3_label += "  ✅ (question generated)"

    with st.expander(step3_label, expanded=st.session_state.exp3_open):
        card(
            "",
            icon="✨",
            body="Now the AI generates a follow-up question and explains why it chose it.",
        )

        resume_text = st.session_state.get("resume_text", "")
        answer_text = st.session_state.get("answer_text", "")
        chosen_scenario = next(
            s for s in SCENARIOS
            if s["name"] == st.session_state.get("scenario_name", SCENARIOS[0]["name"])
        )

        gen = st.button("Generate / Refresh Follow-Up Question", key="btn_generate", type="primary")

        if gen:
            if not (resume_text.strip() and answer_text.strip()):
                st.error("Please make sure both your resume/experience and your answer are filled in.")
            else:
                followup, reasoning, value_tag, confidence, resume_kws, answer_kws = generate_followup(
                    resume_text, answer_text, chosen_scenario["value"]
                )
                st.session_state["followup"] = followup
                st.session_state["reasoning"] = reasoning
                st.session_state["value_tag"] = value_tag
                st.session_state["confidence"] = confidence
                st.session_state["resume_kws"] = resume_kws
                st.session_state["answer_kws"] = answer_kws

                # Generate the follow-up, but keep the participant on Step 3.
                st.session_state.followup_generated = True
                st.session_state.followup_response_done = False
                st.session_state.followup_done = False
                st.session_state.active_step = 3
                st.session_state.exp3_open = True
                st.session_state.exp4_open = False

                st.success("Follow-up generated. Please respond to it below, then continue to Step 4.")
                st.rerun()

        if "followup" in st.session_state:
            st.markdown(
                f"""
                <div style='margin-top:0.5rem; margin-bottom:0.5rem;
                            padding:0.9rem 1rem; border-radius:10px;
                            border-left:4px solid #3b82f6;
                            background-color:#020617;'>
                    <div style='font-size:0.85rem; text-transform:uppercase;
                                letter-spacing:0.05em; color:#9ca3af; margin-bottom:0.25rem;'>
                        AI Follow-Up Question
                    </div>
                    <div style='font-size:1rem; color:#e5e7eb;'>
                        <b>{st.session_state['followup']}</b>
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )

            # Participant responds to the AI follow-up question before moving on
            st.text_area(
                "Your response to the AI follow-up question",
                key="followup_answer_text",
                height=140,
                placeholder="Type your response to the AI follow-up question here…",
            )

            if st.button("Save follow-up response and continue to Step 4", key="btn_followup_answer"):
                if not st.session_state.get("followup_answer_text", "").strip():
                    st.error("Please respond to the AI follow-up question before continuing.")
                else:
                    st.session_state.followup_response_done = True
                    st.session_state.followup_done = True
                    st.session_state.active_step = 4
                    st.session_state.exp3_open = False
                    st.session_state.exp4_open = True
                    st.success("Saved. Moving to Step 4.")
                    st.rerun()


            with st.expander("🧠 Why this question? (click to expand)", expanded=False):
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("**Target value (scenario)**")
                    st.write(st.session_state.get("value_tag", "N/A"))
                    dv, conf = detect_value_tag(answer_text)
                    st.markdown("**System value guess (from your answer)**")
                    st.write(f"{dv} ({conf} confidence)")

                with col2:
                    st.markdown("**Resume keywords detected**")
                    st.write(", ".join(st.session_state.get("resume_kws", [])) or "None detected")

                    st.markdown("**Answer keywords detected**")
                    st.write(", ".join(st.session_state.get("answer_kws", [])) or "None detected")

                st.markdown("**Reasoning summary**")
                st.write(st.session_state.get("reasoning", ""))
else:
    st.info("Complete Step 2 first. Step 3 will appear after you save your scenario + response.")

# ---------------------------------------------------------
# STEP 4 – FAIRNESS & EXPERIENCE FEEDBACK (was Step 5)
# ---------------------------------------------------------
if st.session_state.followup_done:
    step4_label = "Step 4 – Fairness & Experience Feedback"

    with st.expander(step4_label, expanded=st.session_state.exp4_open):
        card(
            "",
            icon="⚖️",
            body="Rate how this follow-up question felt. You can also flag it as unfair or uncomfortable.",
        )

        st.markdown("**Follow-up question shown to you:**")
        st.write(f"_{st.session_state['followup']}_")

        flag_unfair = st.checkbox("I felt this follow-up was unfair, biased, or uncomfortable.")

        unfair_comment = ""
        neutral_q = ""
        if flag_unfair:
            # 1) Start with a softened rephrasing (always begins with the comfort-preface)
            neutral_q = neutralize_question(st.session_state["followup"])
            st.info("In any context you’re comfortable sharing:")

            # 2) Optional: what felt unfair / uncomfortable (about original or rephrased)
            st.text_area(
                "Optional: What felt unfair or uncomfortable ?",
                key="unfair_details",
                height=90,
                placeholder="Example: too personal, unclear, stereotype risk, or not relevant to my response…",
            )
            unfair_comment = st.session_state.get("unfair_details", "")

            # 3) Optional: alternative question (based on the selected scenario)
            st.markdown("**Optional: Alternative question**")
            col_alt_a, col_alt_b = st.columns([1, 3])
            with col_alt_a:
                if st.button("Generate alternative", key="btn_alt_q"):
                    alt = generate_alternative_followup(
                        current_followup=st.session_state.get("followup", ""),
                        value_tag=st.session_state.get("value_tag", "Collaboration"),
                    )
                    st.session_state["alternative_question"] = alt
                    st.session_state["alternative_answer_text"] = ""
                    st.rerun()

            with col_alt_b:
                alt_q = st.session_state.get("alternative_question", "")
                if alt_q:
                    st.write(f"**{alt_q}**")
                else:
                    st.caption("Click “Generate alternative” to view another follow-up question aligned to the same scenario value.")

            st.text_area(
                "Optional: Your response to the alternative question",
                key="alternative_answer_text",
                height=110,
                placeholder="If you prefer the alternative question, you may respond here…",
            )

        st.markdown("#### Quick ratings")
        st.caption("Rating guide: **1 = very low / negative**, **3 = neutral**, **5 = very high / positive**.")

        col_a, col_b = st.columns(2)
        with col_a:
            fairness_score = st.slider(
                "Fairness (Was it unbiased and reasonable?)",
                1, 5, 3,
                help="1=Unfair/bias risk, 3=Neutral, 5=Very fair"
            )
            relevance_score = st.slider(
                "Relevance (Fit your scenario + answer?)",
                1, 5, 3,
                help="1=Not relevant, 3=Somewhat, 5=Highly relevant"
            )
        with col_b:
            comfort_score = st.slider(
                "Comfort (Would you feel okay answering?)",
                1, 5, 3,
                help="1=Very uncomfortable, 3=Neutral, 5=Very comfortable"
            )
            trust_score = st.slider(
                "Trust (Would you trust an interviewer using this AI?)",
                1, 5, 3,
                help="1=No trust, 3=Some trust, 5=High trust"
            )

        accept_ai = st.radio(
            "Would you accept this type of AI interviewer in a real hiring process?",
            ["Yes", "No", "Not sure"],
        )

        open_feedback = st.text_area(
            "Anything else you want to tell us about this question or interface?",
            height=80,
        )

        if st.button("Save and submit feedback", key="btn_save_feedback", type="primary"):
            chosen_scenario = next(
                s for s in SCENARIOS
                if s["name"] == st.session_state.get("scenario_name", SCENARIOS[0]["name"])
            )
            scenario_text = st.session_state.get("scenario_prompt", chosen_scenario["prompt"])

            row = {
                "timestamp": datetime.now().isoformat(),
                "participant_id": st.session_state.get("participant_id", ""),
                "scenario": chosen_scenario["name"],
                "scenario_prompt_used": scenario_text,
                "target_value": chosen_scenario["value"],
                "resume_text": st.session_state.get("resume_text", ""),
                "answer_text": st.session_state.get("answer_text", ""),
                "followup_answer_text": st.session_state.get("followup_answer_text", ""),
                "followup_question": st.session_state.get("followup", ""),
                "reasoning_summary": st.session_state.get("reasoning", ""),
                "resume_keywords": ", ".join(st.session_state.get("resume_kws", [])),
                "answer_keywords": ", ".join(st.session_state.get("answer_kws", [])),
                "value_tag": st.session_state.get("value_tag", ""),
                "confidence": st.session_state.get("confidence", ""),
                "fairness_score": fairness_score,
                "relevance_score": relevance_score,
                "comfort_score": comfort_score,
                "trust_score": trust_score,
                "flag_unfair": flag_unfair,
                "unfair_comment": unfair_comment,
                "alternative_question": st.session_state.get("alternative_question", ""),
                "alternative_answer_text": st.session_state.get("alternative_answer_text", ""),
                "neutralized_question": neutral_q,
                "accept_ai": accept_ai,
                "open_feedback": open_feedback,
            }

            log_row(row)
            st.success("Saved! Your feedback has been added to interview_logs.csv.")

            # --- Downloads ---
            # Best default: CSV (simple + universal) + Excel (for analysis).
            # Word/PDF: best for single-session sharing/appendix.
            if os.path.exists(LOG_FILE):
                df_all = pd.read_csv(LOG_FILE)

                col1, col2, col3, col4 = st.columns(4)

                with col1:
                    with open(LOG_FILE, "rb") as f:
                        st.download_button(
                            "Download CSV (all sessions)",
                            data=f,
                            file_name="interview_logs.csv",
                            mime="text/csv",
                        )

                with col2:
                    st.download_button(
                        "Download Excel (all sessions)",
                        data=df_to_excel_bytes(df_all),
                        file_name="interview_logs.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    )

                with col3:
                    st.download_button(
                        "Download Word (this session)",
                        data=row_to_word_bytes(row),
                        file_name="interview_session_summary.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

                with col4:
                    st.download_button(
                        "Download PDF (this session)",
                        data=row_to_pdf_bytes(row),
                        file_name="interview_session_summary.pdf",
                        mime="application/pdf",
                    )

            st.caption("You can close this window or use the reset button in the sidebar to start again.")

st.caption("Prototype for IS 617 • Human-Centered Computing • Pace University Seidenberg")



